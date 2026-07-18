from pathlib import Path


class CVExtractionError(Exception):
    def __init__(self, code: str, detail: str) -> None:
        self.code = code
        self.detail = detail
        super().__init__(detail)


class CVExtractionService:
    def __init__(self, max_characters: int) -> None:
        self.max_characters = max_characters

    def extract_text(self, path: Path, mime_type: str) -> tuple[str, str, list[str]]:
        if mime_type == "text/plain":
            text = path.read_text(encoding="utf-8", errors="replace")
        elif mime_type == "application/pdf":
            text = self._extract_pdf(path)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = self._extract_docx(path)
        else:
            raise CVExtractionError("UNSUPPORTED_CV_TYPE", "CV file type is not supported.")

        bounded_text = self._bound(text)
        warnings: list[str] = []
        status = "COMPLETED"

        if len(bounded_text.strip()) < 25:
            status = "UNSUPPORTED"
            warnings.append("CV text extraction produced little or no text. Scanned PDFs are not supported in this MVP.")

        return status, bounded_text, warnings

    def _extract_pdf(self, path: Path) -> str:
        try:
            import fitz
        except ImportError as exc:
            raise CVExtractionError("PDF_PARSER_NOT_INSTALLED", "PDF extraction dependency is not installed.") from exc

        chunks: list[str] = []
        with fitz.open(path) as document:
            for page in document:
                chunks.append(page.get_text("text"))

        return "\n\n".join(chunks)

    def _extract_docx(self, path: Path) -> str:
        try:
            import docx
        except ImportError as exc:
            raise CVExtractionError("DOCX_PARSER_NOT_INSTALLED", "DOCX extraction dependency is not installed.") from exc

        document = docx.Document(path)
        chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    chunks.append(" | ".join(cells))

        return "\n".join(chunks)

    def _bound(self, text: str) -> str:
        return text.replace("\x00", "").strip()[: self.max_characters]
